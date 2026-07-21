"""
SafeNet AI - Word Document Generator
=====================================
Generates professional Word documents for URL scan history.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from io import BytesIO


class WordReportGenerator:
    """
    Generate professional Word documents for SafeNet AI scan history.
    """
    
    def __init__(self, user_email, user_plan, scans_data, report_type='all', date_range=None):
        """
        Initialize the Word report generator.
        
        Args:
            user_email (str): User's email address
            user_plan (str): User's current plan
            scans_data (list): List of scan records
            report_type (str): Type of report - 'all', 'phishing', 'legitimate'
            date_range (str): Date range filter
        """
        self.user_email = user_email
        self.user_plan = user_plan
        self.scans_data = scans_data
        self.report_type = report_type
        self.date_range = date_range
        self.buffer = BytesIO()
        
        # Plan display names
        self.plan_display = {
            'free': 'Free',
            'basic': 'Basic',
            'pro': 'Pro',
            'pro_plus': 'Pro Plus',
            'enterprise': 'Enterprise'
        }
        
        # Premium plans
        self.premium_plans = ['pro', 'pro_plus', 'enterprise']
        
    def generate(self):
        """
        Generate the Word document and return as BytesIO buffer.
        
        Returns:
            BytesIO: Word document buffer
        """
        # Create document
        document = Document()
        
        # Set document properties
        document.core_properties.author = 'SafeNet AI'
        document.core_properties.title = 'Scan Security Report'
        
        # ========== HEADER/TITLE ==========
        title = document.add_heading('SafeNet AI', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(124, 58, 237)  # Professional purple
        title_run.font.name = 'Georgia'  # More professional font
        
        subtitle = document.add_heading('Scan Security Report', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
        
        document.add_paragraph()  # Spacer
        
        # ========== USER INFO TABLE ==========
        info_table = document.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        # Populate info table
        info_data = [
            ('User Email:', self.user_email),
            ('Current Plan:', self.plan_display.get(self.user_plan, 'Free')),
            ('Report Type:', self.report_type.capitalize()),
            ('Generated Date:', datetime.utcnow().strftime('%B %d, %Y'))
        ]
        
        for idx, (label, value) in enumerate(info_data):
            row = info_table.rows[idx]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Bold the labels
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        document.add_paragraph()  # Spacer
        
        # ========== SUMMARY SECTION ==========
        document.add_heading('Summary', level=1)
        
        # Calculate statistics
        total_scans = len(self.scans_data)
        phishing_count = sum(1 for scan in self.scans_data 
                           if scan.get('prediction', '').lower() in ['phishing'])
        legitimate_count = sum(1 for scan in self.scans_data 
                             if scan.get('prediction', '').lower() in ['legitimate', 'safe'])
        suspicious_count = sum(1 for scan in self.scans_data 
                             if scan.get('prediction', '').lower() in ['suspicious'])
        
        # Add summary as bullet points
        summary_para = document.add_paragraph()
        summary_para.add_run(f'Total URLs Scanned: ').bold = True
        summary_para.add_run(str(total_scans))
        
        summary_para = document.add_paragraph()
        summary_para.add_run(f'Phishing Detected: ').bold = True
        summary_para.add_run(str(phishing_count))
        
        summary_para = document.add_paragraph()
        summary_para.add_run(f'Legitimate URLs: ').bold = True
        summary_para.add_run(str(legitimate_count))
        
        summary_para = document.add_paragraph()
        summary_para.add_run(f'Suspicious URLs: ').bold = True
        summary_para.add_run(str(suspicious_count))
        
        summary_para = document.add_paragraph()
        summary_para.add_run(f'Detection Method: ').bold = True
        summary_para.add_run('Results based on AI-powered multi-layer detection system')
        
        document.add_paragraph()  # Spacer
        
        # ========== SCAN HISTORY TABLE ==========
        document.add_heading('Scan History', level=1)
        
        # Check if free user and limit scans
        is_free_user = self.user_plan == 'free'
        scans_to_show = self.scans_data[:10] if is_free_user else self.scans_data
        
        if is_free_user and total_scans > 10:
            # Add upgrade message
            warning = document.add_paragraph()
            warning_run = warning.add_run(
                f'⚠️ Free plan: Showing only 10 of {total_scans} scans. '
                f'Upgrade to Pro or Pro+ to export all scans.'
            )
            warning_run.font.color.rgb = RGBColor(255, 107, 107)  # Red
            warning_run.font.italic = True
        
        # Create table with headers
        table = document.add_table(rows=1, cols=5)
        table.style = 'Medium Shading 1 Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['#', 'URL', 'Result', 'Risk Level', 'Scan Date']
        for idx, header in enumerate(headers):
            header_cells[idx].text = header
            header_cells[idx].paragraphs[0].runs[0].font.bold = True
        
        # Add data rows
        for idx, scan in enumerate(scans_to_show, 1):
            row = table.add_row().cells
            
            url = scan.get('url', 'N/A')
            # Truncate long URLs
            if len(url) > 60:
                url = url[:57] + '...'
            
            result = scan.get('prediction', 'Unknown')
            
            # Determine risk level
            if result.lower() in ['phishing']:
                risk_level = 'High'
            elif result.lower() in ['suspicious']:
                risk_level = 'Medium'
            else:
                risk_level = 'Low'
            
            # Format date
            timestamp = scan.get('timestamp')
            if timestamp:
                if isinstance(timestamp, str):
                    scan_date = timestamp
                else:
                    scan_date = timestamp.strftime('%b %d, %Y')
            else:
                scan_date = 'N/A'
            
            # Populate row
            row[0].text = str(idx)
            row[1].text = url
            row[2].text = result
            row[3].text = risk_level
            row[4].text = scan_date
            
            # Color code results
            result_run = row[2].paragraphs[0].runs[0]
            if result.lower() == 'phishing':
                result_run.font.color.rgb = RGBColor(231, 76, 60)  # Red
                result_run.font.bold = True
            elif result.lower() == 'suspicious':
                result_run.font.color.rgb = RGBColor(243, 156, 18)  # Orange
            elif result.lower() in ['legitimate', 'safe']:
                result_run.font.color.rgb = RGBColor(39, 174, 96)  # Green
        
        document.add_paragraph()  # Spacer
        
        # ========== FOOTER ==========
        footer_para = document.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            'Generated by SafeNet AI – AI-Powered URL Security'
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(153, 153, 153)
        
        # Add premium watermark
        if self.user_plan in self.premium_plans:
            watermark_para = document.add_paragraph()
            watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            watermark_run = watermark_para.add_run(
                f'✓ {self.plan_display[self.user_plan]} Feature'
            )
            watermark_run.font.size = Pt(8)
            watermark_run.font.color.rgb = RGBColor(52, 152, 219)  # Blue
        
        # Save to buffer
        document.save(self.buffer)
        self.buffer.seek(0)
        return self.buffer


def generate_word_report(user_email, user_plan, scans_data, report_type='all', date_range=None):
    """
    Helper function to generate a Word report.
    
    Args:
        user_email (str): User's email
        user_plan (str): User's plan
        scans_data (list): List of scan records
        report_type (str): 'all', 'phishing', or 'legitimate'
        date_range (str): '7', '30', 'custom', or None
        
    Returns:
        BytesIO: Word document buffer
    """
    generator = WordReportGenerator(user_email, user_plan, scans_data, report_type, date_range)
    return generator.generate()
